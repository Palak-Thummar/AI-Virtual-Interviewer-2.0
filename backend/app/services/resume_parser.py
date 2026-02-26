"""
Resume parsing service for PDF and DOCX files.
Extracts text content from resume documents.
"""

import pdfplumber
from docx import Document
from typing import Tuple
import os


async def extract_resume_text(file_path: str, file_type: str) -> str:
    """
    Extract text from resume file (PDF or DOCX).
    
    Args:
        file_path: Path to the resume file
        file_type: Type of file ('pdf' or 'docx')
        
    Returns:
        Extracted text from resume
        
    Raises:
        ValueError: If file type is not supported or extraction fails
    """
    try:
        if file_type.lower() == "pdf":
            return await _extract_pdf_text(file_path)
        elif file_type.lower() == "docx":
            return await _extract_docx_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    except Exception as e:
        raise ValueError(f"Failed to extract resume text: {str(e)}")


async def _extract_pdf_text(file_path: str) -> str:
    """Extract text from PDF file."""
    text_content = []
    
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
    except Exception as e:
        raise ValueError(f"PDF extraction failed: {str(e)}")
    
    return "\n".join(text_content)


async def _extract_docx_text(file_path: str) -> str:
    """Extract text from DOCX file."""
    text_content = []
    
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_content.append(cell.text)
    except Exception as e:
        raise ValueError(f"DOCX extraction failed: {str(e)}")
    
    return "\n".join(text_content)


def clean_text(text: str) -> str:
    """
    Clean and normalize extracted text.
    
    Args:
        text: Raw text to clean
        
    Returns:
        Cleaned text
    """
    # Remove extra whitespace
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    return "\n".join(lines)
